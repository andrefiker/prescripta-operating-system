from __future__ import annotations

import json
from pathlib import Path
from textwrap import dedent

from docx import Document
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / 'public' / 'documents'

PROFILE = {
    'legal_name': 'Prescripta Tecnologia em Saude Ltda.',
    'brand': 'Prescripta',
    'city': 'Sao Paulo, SP, Brasil',
    'website': 'https://prescriptamed.com.br/',
    'privacy_email': 'privacidade@prescripta.com.br',
    'support_email': 'suporte@prescripta.com.br',
    'commercial_email': 'comercial@prescripta.com.br',
    'primary_offer': 'infraestrutura operacional para recorrencia prescricional em psiquiatria privada',
    'solo_price': 'R$ 349/mes',
    'clinic_price': 'R$ 2.490/mes',
    'pilot_policy': 'piloto pago de 60 dias com metas de ativacao, validacao operacional e decisao formal de conversao',
}


def make_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name='BodyCompact',
            parent=styles['BodyText'],
            fontName='Helvetica',
            fontSize=9.5,
            leading=13,
            textColor=colors.HexColor('#2b2117'),
            alignment=TA_LEFT,
            spaceAfter=6,
        )
    )
    styles['Heading1'].fontName = 'Helvetica-Bold'
    styles['Heading1'].fontSize = 20
    styles['Heading1'].leading = 24
    styles['Heading1'].textColor = colors.HexColor('#8b4f1f')
    styles['Heading2'].fontName = 'Helvetica-Bold'
    styles['Heading2'].fontSize = 12
    styles['Heading2'].leading = 15
    styles['Heading2'].spaceBefore = 10
    styles['Heading2'].textColor = colors.HexColor('#2b2117')
    return styles


def markdown_to_story(markdown: str):
    styles = make_styles()
    story = []
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            story.append(Spacer(1, 4))
            continue
        if line.startswith('# '):
            story.append(Paragraph(line[2:], styles['Heading1']))
            story.append(Spacer(1, 4))
            continue
        if line.startswith('## '):
            story.append(Paragraph(line[3:], styles['Heading2']))
            continue
        if line.startswith('### '):
            story.append(Paragraph(f'<b>{line[4:]}</b>', styles['BodyCompact']))
            continue
        if line.startswith('- '):
            story.append(Paragraph(f'&bull; {line[2:]}', styles['BodyCompact']))
            continue
        story.append(Paragraph(line, styles['BodyCompact']))
    return story


def save_pdf(path: Path, markdown: str):
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    doc.build(markdown_to_story(markdown))


def save_docx(path: Path, markdown: str):
    document = Document()
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            document.add_paragraph('')
            continue
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
            continue
        if line.startswith('## '):
            document.add_heading(line[3:], level=2)
            continue
        if line.startswith('### '):
            document.add_heading(line[4:], level=3)
            continue
        if line.startswith('- '):
            document.add_paragraph(line[2:], style='List Bullet')
            continue
        document.add_paragraph(line)
    document.save(path)


def save_revenue_xlsx(path: Path):
    wb = Workbook()
    assumptions = wb.active
    assumptions.title = 'Assumptions'
    assumptions.append(['Metric', 'Value'])
    assumptions.append(['Solo monthly price', 349])
    assumptions.append(['Clinic monthly price', 2490])
    assumptions.append(['Solo new logos per month', 4])
    assumptions.append(['Clinic new logos per month', 1])
    assumptions.append(['Solo monthly churn', 0.03])
    assumptions.append(['Clinic monthly churn', 0.015])
    assumptions.append(['Fixed monthly costs', 35000])

    monthly = wb.create_sheet('MRR Projection')
    monthly.append(
        [
            'Month',
            'Starting Solo',
            'New Solo',
            'Churned Solo',
            'Ending Solo',
            'Starting Clinic',
            'New Clinic',
            'Churned Clinic',
            'Ending Clinic',
            'Solo MRR',
            'Clinic MRR',
            'Total MRR',
            'Fixed Costs',
            'Contribution Before Variable Costs',
        ]
    )

    ending_solo = 0.0
    ending_clinic = 0.0
    for month in range(1, 13):
        starting_solo = ending_solo
        starting_clinic = ending_clinic
        new_solo = 4
        new_clinic = 1
        churned_solo = round(starting_solo * 0.03, 2)
        churned_clinic = round(starting_clinic * 0.015, 2)
        ending_solo = starting_solo + new_solo - churned_solo
        ending_clinic = starting_clinic + new_clinic - churned_clinic
        solo_mrr = round(ending_solo * 349, 2)
        clinic_mrr = round(ending_clinic * 2490, 2)
        total_mrr = solo_mrr + clinic_mrr
        monthly.append(
            [
                month,
                starting_solo,
                new_solo,
                churned_solo,
                ending_solo,
                starting_clinic,
                new_clinic,
                churned_clinic,
                ending_clinic,
                solo_mrr,
                clinic_mrr,
                total_mrr,
                35000,
                total_mrr - 35000,
            ]
        )

    wb.save(path)


def base_assets():
    p = PROFILE
    return [
        {
            'key': 'termos-de-uso',
            'title': 'Termos de Uso',
            'tier': 'Tier 1',
            'summary': 'Termos de uso completos para SaaS de saude.',
            'pdf': True,
            'markdown': dedent(
                f"""
                # Termos de Uso - {p['brand']}

                ## 1. Partes e objeto
                Estes Termos de Uso regulam o acesso e a utilizacao da plataforma {p['brand']}, operada por {p['legal_name']}, com sede em {p['city']}. A plataforma e disponibilizada em modelo SaaS B2B para apoiar a organizacao operacional da recorrencia prescricional, incluindo rotinas de prazo, retorno, renovacao, confirmacao, fila operacional e trilha de atividades.

                ## 2. Natureza da plataforma
                A {p['brand']} nao atua como prontuario eletronico generico, nao substitui a autonomia tecnica do profissional de saude e nao realiza decisao clinica autonoma. Toda decisao medica, prescricional, diagnostica ou terapeutica permanece sob responsabilidade do profissional habilitado e da instituicao contratante.

                ## 3. Licenca e conta
                A contratante recebe licenca limitada, nao exclusiva, nao transferivel e revogavel para uso da plataforma durante a vigencia contratual. O acesso e restrito a pessoas autorizadas, com credenciais individuais e intransferiveis. E vedado copiar, revender, sublicenciar, descompilar ou explorar comercialmente a plataforma fora do contrato.

                ## 4. Escopo e limites
                O escopo contratado sera definido na proposta comercial, ordem de servico ou contrato SaaS aplicavel. Itens em roadmap, especialmente os dependentes de validacao regulatoria especifica, nao integram a entrega padrao ate confirmacao formal da {p['brand']}. A plataforma nao dispensa avaliacao juridica, regulatoria, clinica ou operacional propria da contratante.

                ## 5. Responsabilidades da contratante
                A contratante deve manter base legal apropriada para os tratamentos que lhe competem, definir processos internos de uso, revisar toda decisao clinica, nomear owner operacional e garantir que seus profissionais atuem dentro das normas eticas e regulatorias aplicaveis.

                ## 6. Responsabilidades da {p['brand']}
                A {p['brand']} devera fornecer a plataforma nos limites do escopo contratado, manter canais de suporte conforme o plano, adotar medidas tecnicas e organizacionais compativeis com o risco e disponibilizar documentacao basica de privacidade, seguranca, subprocessadores e politica de IA.

                ## 7. Dados pessoais e seguranca
                A plataforma pode tratar dados pessoais e dados pessoais sensiveis, inclusive dados de saude, em nome da contratante e conforme o fluxo contratado. O tratamento deve observar a LGPD, os principios de necessidade, finalidade, adequacao, seguranca e responsabilizacao. A {p['brand']} podera manter logs, metadados e trilhas tecnicas para seguranca, auditoria, suporte e resguardo contratual.

                ## 8. Disponibilidade, cobranca e encerramento
                A {p['brand']} buscara disponibilidade comercialmente razoavel, admitidas manutencoes programadas, eventos de seguranca, falhas de terceiros e hipoteses de forca maior. Valores, setup, reajustes, renovacao, inadimplencia, suspensao e cancelamento seguirao o instrumento comercial aplicavel. A inadimplencia pode ensejar suspensao de acesso.

                ## 9. Propriedade intelectual e confidencialidade
                Todos os direitos relativos a software, arquitetura, documentacao, layout, marcas, processos e melhorias da plataforma pertencem a {p['legal_name']} ou a seus licenciantes. As partes devem preservar a confidencialidade das informacoes tecnicas, comerciais, operacionais e de seguranca.

                ## 10. Contato, foro e aviso
                Duvidas podem ser encaminhadas para {p['support_email']}. Fica eleito o foro da comarca de Sao Paulo, sem prejuizo de disposicao contratual especifica, para dirimir controversias oriundas destes Termos. Esta minuta deve ser revisada por assessoria juridica antes do uso final.
                """
            ).strip(),
        },
        {
            'key': 'politica-de-privacidade',
            'title': 'Politica de Privacidade',
            'tier': 'Tier 1',
            'summary': 'Politica completa com LGPD e direitos do titular.',
            'pdf': True,
            'markdown': dedent(
                f"""
                # Politica de Privacidade - {p['brand']}

                ## 1. Quem somos
                Esta Politica de Privacidade descreve como {p['legal_name']} trata dados pessoais no contexto de suas atividades corporativas, comerciais, contratuais, operacionais e de disponibilizacao da plataforma {p['brand']}.

                ## 2. Escopo, categorias e finalidades
                Esta politica cobre site, relacionamento comercial, onboarding, suporte, faturamento, seguranca e uso da plataforma. Podemos tratar dados cadastrais, profissionais, de contato, autenticacao, navegacao, logs, faturamento, suporte, relacao contratual e, quando aplicavel ao servico, dados pessoais sensiveis relacionados a saude e rotina operacional de prescricoes recorrentes. As finalidades incluem negociacao, contratacao, execucao do servico, seguranca, exercicio regular de direitos, cumprimento de obrigacoes legais, melhoria operacional e suporte.

                ## 3. Bases legais
                As bases legais podem incluir execucao de contrato, cumprimento de obrigacao legal ou regulatoria, exercicio regular de direitos, legitimo interesse, tutela da saude quando aplicavel e outras hipoteses previstas na LGPD. O enquadramento concreto depende da atividade de tratamento.

                ## 4. Papel da {p['brand']}
                Em marketing, relacao comercial, faturamento, gestao contratual e operacao corporativa propria, a {p['brand']} tende a atuar como controladora. Em funcionalidades executadas em nome de clientes, relacionadas ao uso contratado da plataforma, a {p['brand']} tende a atuar como operadora, conforme o contrato e o DPA aplicavel.

                ## 5. Compartilhamento, subprocessadores e transferencias
                Dados podem ser compartilhados com prestadores de servico, subprocessadores, fornecedores de infraestrutura, ferramentas de atendimento, plataformas de email, parceiros de analytics, consultores juridicos e autoridades publicas quando necessario. Eventuais transferencias internacionais observarao salvaguardas razoaveis cabiveis.

                ## 6. Retencao, descarte e seguranca
                Os dados serao mantidos pelo tempo necessario ao cumprimento das finalidades declaradas, do contrato, da legislacao aplicavel e do exercicio regular de direitos. Encerrado o tratamento ou expirado o prazo de retencao, os dados serao eliminados, anonimizados ou arquivados conforme criterio legal, contratual e de seguranca. Adotamos medidas tecnicas e organizacionais proporcionais ao risco, incluindo gestao de acesso, logs, procedimentos internos, avaliacao de fornecedores e resposta a incidentes.

                ## 7. Direitos dos titulares
                O titular pode solicitar confirmacao de tratamento, acesso, correcao, anonimização, bloqueio, eliminacao, portabilidade quando cabivel, informacoes sobre compartilhamento e demais direitos previstos na LGPD, observadas limitacoes legais e tecnicas aplicaveis. Solicitacoes podem ser enviadas para {p['privacy_email']}.

                ## 8. Atualizacoes e contato
                Esta politica pode ser alterada a qualquer tempo para refletir evolucao legal, regulatoria, contratual ou operacional. A versao vigente sera publicada em {p['website']}. Canal de privacidade: {p['privacy_email']}. Canal geral de suporte: {p['support_email']}. Esta minuta deve ser validada por assessoria juridica antes da publicacao final.
                """
            ).strip(),
        },
        {
            'key': 'dpa',
            'title': 'DPA',
            'tier': 'Tier 1',
            'summary': 'Anexo contratual de protecao de dados.',
            'pdf': True,
            'markdown': dedent(
                f"""
                # DPA - Data Processing Agreement - {p['brand']}

                ## 1. Finalidade e escopo
                Este DPA integra a proposta comercial, ordem de servico ou contrato SaaS celebrado entre {p['legal_name']} e a contratante. Ele regula o tratamento de dados pessoais realizado pela {p['brand']} no contexto da prestacao de servicos SaaS, especialmente quando a {p['brand']} atuar como operadora em nome da contratante.

                ## 2. Instrucao documentada e finalidades autorizadas
                A {p['brand']} tratara dados pessoais de acordo com as instrucoes documentadas da contratante, com o escopo contratado e com a legislacao aplicavel. O tratamento inclui hospedagem, organizacao, exibicao, manutencao, suporte, trilhas de uso, autenticacao, seguranca, backup, implantacao e demais operacoes necessarias para a execucao do servico contratado.

                ## 3. Categorias de dados, seguranca e confidencialidade
                As categorias de dados e titulares dependerao do uso concreto contratado, podendo abranger profissionais de saude, equipe operacional, representantes da contratante e pacientes envolvidos no fluxo de recorrencia prescricional. A {p['brand']} adotara medidas tecnicas e organizacionais proporcionais ao risco e assegurara que pessoas autorizadas a tratar dados estejam sujeitas a obrigacoes de confidencialidade compatveis com a natureza dos dados tratados.

                ## 4. Subprocessadores, cooperacao e incidentes
                A {p['brand']} podera utilizar subprocessadores para hospedar, operar, monitorar, comunicar, atender ou prestar suporte ao servico, desde que imponha obrigacoes adequadas de protecao de dados e mantenha lista atualizada dos subprocessadores relevantes. A {p['brand']} prestara assistencia proporcional para atendimento a direitos dos titulares e comunicara a contratante sem atraso indevido quando tomar conhecimento de incidente de seguranca relevante no contexto do servico.

                ## 5. Retencao, exclusao e auditoria documental
                Encerrado o contrato, a {p['brand']} observara as regras contratuais de retencao, exportacao quando aplicavel, descarte seguro e arquivamento minimo para exercicio regular de direitos e obrigacoes legais ou reguladoras. Mediante pedido razoavel, a {p['brand']} podera disponibilizar informacoes documentais sobre praticas de seguranca, privacidade, subprocessadores e politicas aplicaveis, resguardados segredos comerciais e proporcionalidade operacional. Esta versao e uma minuta-base e deve ser revista pelo juridico antes do uso final.
                """
            ).strip(),
        },
        {
            'key': 'politica-de-ia',
            'title': 'Politica de IA',
            'tier': 'Tier 1',
            'summary': 'Politica completa sobre automacao e supervisao humana.',
            'pdf': True,
            'markdown': dedent(
                f"""
                # Politica de IA - {p['brand']}

                ## 1. Objetivo e escopo
                Esta politica estabelece os principios de uso de recursos de automacao e IA no contexto da plataforma {p['brand']}. Ela cobre recursos de classificacao, organizacao, apoio a fluxo, resumo operacional, sugestao de proximo passo, ordenacao de fila, ajuda contextual e automacoes similares.

                ## 2. Principio fundamental
                A {p['brand']} nao realiza decisao clinica autonoma. Nenhum recurso de IA substitui julgamento medico, ato prescricional, diagnostico, definicao terapeutica ou responsabilidade profissional da contratante.

                ## 3. Supervisao humana e limites
                Toda funcionalidade assistida deve operar sob supervisao humana apropriada ao contexto. Outputs relevantes para uso clinico ou operacional critico devem ser revisados por usuario habilitado antes de gerar efeito externo, decisao sensivel ou acao prescritiva. Sao vedados claims publicos que superdimensionem a autonomia do sistema.

                ## 4. Qualidade, dados e governanca
                Saidas assistidas por software podem conter imprecisao, omissao ou classificacao inadequada. Usuarios devem adotar postura de revisao critica. O uso de dados em fluxos assistidos deve observar necessidade, adequacao e minimizacao. Revisoes de risco e aprovacoes de claims devem envolver lideranca executiva, produto e juridico externo quando necessario.

                ## 5. Transparencia e contato
                A existencia de automacao relevante deve ser comunicada de maneira clara nos materiais de produto, trust center e contratos quando aplicavel. Duvidas sobre esta politica podem ser enviadas para {p['privacy_email']} ou {p['support_email']}.
                """
            ).strip(),
        },
        {
            'key': 'proposta-comercial',
            'title': 'Proposta comercial completa',
            'tier': 'Tier 1',
            'summary': 'Minuta completa com campos, escopo, pricing e cronograma.',
            'pdf': True,
            'docx': True,
            'markdown': dedent(
                f"""
                # Proposta Comercial Completa - {p['brand']}

                ## 1. Dados da oportunidade
                - Cliente:
                - CNPJ ou CPF:
                - Responsavel:
                - Email:
                - Data da proposta:
                - Validade:

                ## 2. Resumo executivo
                A {p['brand']} e uma plataforma SaaS B2B focada em {p['primary_offer']}. O objetivo desta proposta e estruturar o fluxo operacional entre termino previsto, retorno, renovacao, confirmacao e acompanhamento da recorrencia prescricional, com supervisao humana e governanca adequada ao contexto de saude.

                ## 3. Dor observada no discovery
                Preencher com o problema dominante da conta:
                - [ ] Retrabalho entre agenda, WhatsApp e planilhas
                - [ ] Falta de visibilidade sobre prazos e pendencias
                - [ ] Baixa continuidade entre consulta e renovacao
                - [ ] Falta de fila central e ownership operacional
                - [ ] Risco de crescimento sem processo

                ## 4. Escopo contratado
                Esta proposta contempla licenca de uso da plataforma, configuracao operacional inicial, treinamento por perfil, suporte conforme plano e revisao de ativacao. Nao fazem parte do escopo padrao prontuario generico, decisao clinica automatizada, integracoes nao aprovadas e funcionalidades em roadmap regulatorio sem formalizacao.

                ## 5. Planos e investimento
                | Plano | Perfil | Valor recorrente | Setup |
                | --- | --- | --- | --- |
                | Start Solo | psiquiatra individual | {p['solo_price']} | incluido |
                | Clinica | estrutura multiprofissional | {p['clinic_price']} | sob definicao em proposta |
                | Enterprise | rede, grupo ou parceiro | sob proposta | sob proposta |

                ## 6. Premissas e cronograma
                A contratacao segue o escopo e limite do plano aprovado. Requisitos especiais de seguranca, SLA, SSO, rollout por unidade, customizacao ou integracao serao objeto de escopo adicional. Cronograma sugerido: semana 1 assinatura e kickoff; semana 2 configuracao; semana 3 treinamento; semana 4 go-live; semana 6 revisao de ativacao.

                ## 7. Base juridica e aceite
                O relacionamento sera regulado pelo contrato SaaS, DPA e politicas vigentes da {p['brand']}. Nome: Cargo: Assinatura: Data:. Esta proposta e uma minuta comercial base e deve ser revisada antes do envio final a cada cliente.
                """
            ).strip(),
        },
        {
            'key': 'one-pager',
            'title': 'One-pager',
            'tier': 'Tier 1',
            'summary': 'Resumo executivo pronto para envio apos demo.',
            'pdf': True,
            'markdown': dedent(
                f"""
                # One-pager - {p['brand']}

                ## Headline
                A camada operacional da recorrencia prescricional para psiquiatria privada.

                ## Subheadline
                Menos improviso entre consulta, prazo e renovacao. Mais visibilidade, ownership e continuidade operacional.

                ## O problema
                Consultorios e clinicas ainda dependem de agenda, WhatsApp, memoria individual e retrabalho manual para acompanhar termino previsto, retorno e renovacao. Isso reduz previsibilidade, aumenta carga cognitiva e cria risco operacional.

                ## O que a {p['brand']} faz
                A {p['brand']} organiza a rotina de recorrencia prescricional em um fluxo observavel, com fila, status, contexto, alerta e acompanhamento assistido. A plataforma entra de forma estreita no processo em que a dor realmente acontece.

                ## O que nao faz
                - Nao e prontuario generico.
                - Nao toma decisao clinica autonoma.
                - Nao promete cobertura regulatoria fora do escopo contratado.

                ## Confianca e CTA
                LGPD, DPA, politica de IA com supervisao humana, trilha operacional e direcao clara sobre limites de escopo e roadmap regulatorio. Agende uma demo diagnostica ou solicite um piloto pago estruturado em {p['commercial_email']}.
                """
            ).strip(),
        },
        {
            'key': 'contrato-saas-msa',
            'title': 'Contrato SaaS (MSA)',
            'tier': 'Tier 1',
            'summary': 'Contrato-base para assinatura com clientes.',
            'pdf': True,
            'markdown': dedent(
                f"""
                # Contrato SaaS (MSA) - {p['brand']}

                ## 1. Partes e objeto
                De um lado, {p['legal_name']}; de outro, a contratante identificada na proposta ou ordem de servico aplicavel. O presente contrato regula o licenciamento de uso da plataforma {p['brand']}, bem como os servicos acessorios de implantacao, treinamento e suporte descritos no instrumento comercial aplicavel.

                ## 2. Escopo, vigencia e pagamento
                O plano, usuarios, servicos incluidos, setup, valores, prazos, funcionalidades, exclusoes, SLA e anexos aplicaveis constarao da proposta correspondente. A vigencia inicial e a renovacao seguirao o instrumento comercial. A contratante pagara os valores recorrentes e eventuais fees de setup, implantacao, customizacao ou servicos adicionais conforme proposta aprovada.

                ## 3. Obrigacoes da {p['brand']}
                Disponibilizar a plataforma dentro do escopo contratado, executar onboarding, suporte e governanca minima nos termos do plano, manter documentacao basica de privacidade, seguranca, IA e subprocessadores e atuar com medidas tecnicas e organizacionais proporcionais ao risco.

                ## 4. Obrigacoes da contratante
                Utilizar a plataforma dentro do escopo legal e etico aplicavel, nomear responsavel operacional, garantir que seus usuarios estejam autorizados, manter a revisao e responsabilidade sobre qualquer decisao clinica e cumprir as bases legais e obrigacoes reguladoras que lhe competem.

                ## 5. Dados, seguranca e propriedade intelectual
                As partes observarao a LGPD e demais regras aplicaveis. O tratamento de dados pessoais sera disciplinado por este contrato e por DPA especifico. Todo o software, documentacao, metodologia, marca, estrutura tecnica e evolucoes da plataforma sao e permanecem de titularidade da {p['legal_name']} ou de seus licenciantes.

                ## 6. Confidencialidade, responsabilidade e rescisao
                As partes obrigam-se a preservar confidencialidade sobre informacoes tecnicas, comerciais, operacionais, contratuais, de seguranca e de dados. A {p['brand']} nao se responsabiliza por decisoes clinicas, uso indevido da plataforma, descumprimento legal da contratante ou falhas de terceiros alheios ao seu controle. O contrato podera ser rescindido por descumprimento relevante, inadimplencia, risco legal relevante, encerramento nos termos comerciais acordados ou comum acordo. Esta minuta deve ser revisada por assessoria juridica antes da assinatura com clientes.
                """
            ).strip(),
        },
        {
            'key': 'playbook-onboarding',
            'title': 'Playbook de onboarding',
            'tier': 'Tier 2',
            'summary': 'Checklist, treino por papel e go-live.',
            'pdf': True,
            'markdown': dedent(
                """
                # Playbook de Onboarding - Prescripta

                ## 1. Objetivo
                Colocar a conta em operacao previsivel em ate 30 dias, com ownership claro, usuarios ativos e rotina principal executada sem improviso.

                ## 2. D1 - handoff comercial
                Confirmar escopo vendido, plano, setup, numero de usuarios e promessa feita. Registrar dor dominante, urgencia e metrica de sucesso. Nomear owner interno da Prescripta e owner da conta no cliente. Garantir envio de proposta assinada, contrato e DPA.

                ## 3. Kickoff
                Relembrar objetivo de negocio da conta. Mapear fluxo atual entre consulta, termino previsto, retorno, renovacao e confirmacao. Identificar pessoas por papel: prescritor, secretaria, recepcao, administrador. Fechar cronograma de configuracao, treinamento e go-live.

                ## 4. Script de treinamento por papel
                ### Prescritor
                Mostrar a tese da plataforma em 2 minutos, navegar pela rotina principal, mostrar onde entra revisao humana e validar como aprovar ou encaminhar pendencias.
                ### Secretaria / recepcao
                Explicar ownership da fila, demonstrar status e proximo passo, treinar confirmacao, pendencia e escalonamento e simular um caso de atraso.
                ### Administrador
                Explicar usuarios, permissoes, governanca, relatorios, canais de suporte e revisao de ativacao.

                ## 5. Checklists
                - Usuarios convidados e ativos
                - Fluxo principal configurado
                - Contrato e DPA em ordem
                - Responsavel do cliente disponivel
                - Primeira revisao agendada

                ## 6. Revisao de ativacao
                Perguntar se a rotina principal esta sendo usada, onde o fluxo ainda quebra, qual papel esta travando a operacao, se o valor percebido ja apareceu e se a conta esta pronta para operar sem suporte intenso.
                """
            ).strip(),
        },
        {
            'key': 'discovery-call-script',
            'title': 'Discovery call script',
            'tier': 'Tier 2',
            'summary': 'Script completo com logica de qualificacao.',
            'pdf': True,
            'markdown': dedent(
                """
                # Discovery Call Script - Prescripta

                ## Objetivo
                Qualificar dor, buyer, urgencia, formato de operacao e chance de ativacao antes de demo ou proposta.

                ## Bloco 1 - contexto
                Quantos prescritores participam da rotina? Quem opera o fluxo no dia a dia? Quais ferramentas estao envolvidas hoje? Qual o volume aproximado de pacientes ativos em recorrencia?

                ## Bloco 2 - dor
                Onde a operacao quebra mais: prazo, contexto, follow-up, confirmacao ou ownership? O que acontece quando um paciente some entre consulta e renovacao? Quais tarefas hoje dependem de memoria, agenda ou WhatsApp?

                ## Bloco 3 - urgencia e autoridade
                O problema atrapalha receita, rotina, equipe ou risco? Existe pressao para resolver isso agora? Quem decide comprar? Quem usaria diariamente? Quem precisa aprovar juridico, financeiro ou operacao?

                ## Qualificacao
                Criticos de qualificacao: dor repetitiva e real, owner operacional identificado, capacidade de pagar, disposicao para onboarding e escopo aderente ao produto atual.

                ## Logica de aprofundamento
                Se disser que ja faz isso no WhatsApp, perguntar como visibilidade, ownership e historico sao mantidos. Se pedir IA autonoma, reposicionar a Prescripta como infraestrutura operacional com supervisao humana. Se nao houver owner operacional, registrar risco alto de ativacao e evitar proposta imediata.
                """
            ).strip(),
        },
        {
            'key': 'welcome-email-sequence',
            'title': 'Welcome email sequence',
            'tier': 'Tier 2',
            'summary': 'Sequencia pronta do signup a ativacao.',
            'pdf': True,
            'markdown': dedent(
                f"""
                # Welcome Email Sequence - {p['brand']}

                ## Email 1 - assinatura confirmada
                Assunto: Bem-vindo(a) a {p['brand']}
                Corpo: Seu onboarding com a {p['brand']} comeca agora. Nosso objetivo nos proximos dias e colocar sua rotina principal em operacao com ownership claro, menos improviso e uma revisao formal de ativacao.

                ## Email 2 - preparacao para kickoff
                Assunto: O que precisamos antes do kickoff
                Corpo: Precisamos de owner operacional nomeado, lista inicial de usuarios e descricao breve do fluxo atual entre consulta, prazo, retorno e renovacao.

                ## Email 3 - treinamento agendado
                Assunto: Treinamento da equipe e proximo passo
                Corpo: O foco do treinamento nao e aprender tela por tela, e sim operar a rotina real da sua conta.

                ## Email 4 - go-live
                Assunto: Sua conta entra em go-live
                Corpo: A partir de hoje, a recomendacao e operar sua rotina principal pela {p['brand']}. Se algo travar, nosso suporte esta disponivel em {p['support_email']}.

                ## Email 5 - primeira revisao
                Assunto: Revisao da primeira semana
                Corpo: Queremos revisar o que ja entrou em rotina, quais pontos ainda geram atrito e o que precisa de ajuste fino.

                ## Email 6 - ativacao
                Assunto: Fechando a etapa de ativacao
                Corpo: Se os usuarios criticos ja estao ativos e a rotina principal esta rodando, entramos na fase de operacao recorrente.
                """
            ).strip(),
        },
        {
            'key': 'trust-center-content',
            'title': 'Trust center page content',
            'tier': 'Tier 2',
            'summary': 'Copy pronta das secoes do trust center.',
            'pdf': True,
            'markdown': dedent(
                f"""
                # Trust Center - Copy Base - {p['brand']}

                ## Hero
                A {p['brand']} foi desenhada para operar uma camada estreita e critica da recorrencia prescricional com governanca, supervisao humana e clareza de limites.

                ## Privacidade
                Tratamos dados pessoais em contextos distintos, ora como controladora em nossas atividades proprias, ora como operadora em nome de clientes. Mantemos politica de privacidade, canal para direitos dos titulares e DPA comercial para apoiar o fechamento de contratos.

                ## Seguranca
                Mantemos medidas tecnicas e organizacionais proporcionais ao risco, incluindo controle de acesso, logs, procedimentos internos, governanca de fornecedores e resposta a incidentes.

                ## IA e automacao
                A {p['brand']} nao toma decisao clinica autonoma. Qualquer recurso assistido por software ou IA e desenhado para apoiar contexto, fila e operacao, sempre com revisao humana apropriada ao uso.

                ## Subprocessadores e incidentes
                Utilizamos provedores de infraestrutura e servicos de apoio que podem atuar como subprocessadores. Mantemos processo interno para triagem, resposta e comunicacao de incidentes relevantes.

                ## Contatos
                Privacidade: {p['privacy_email']}
                Suporte: {p['support_email']}
                Comercial: {p['commercial_email']}
                """
            ).strip(),
        },
        {
            'key': 'crm-setup-spec',
            'title': 'CRM setup spec',
            'tier': 'Tier 3',
            'summary': 'Campos, estagios, automacoes e relatorios.',
            'pdf': True,
            'markdown': dedent(
                """
                # CRM Setup Spec - Prescripta

                ## Ferramenta recomendada
                Pipedrive ou HubSpot free/Starter para pipeline principal. Se a equipe ainda for muito pequena, usar planilha mestre apenas como espelho de controle, nunca como unica fonte.

                ## Estagios
                1. Lead nomeado
                2. Discovery agendado
                3. Discovery realizado
                4. Demo agendada
                5. Demo realizada
                6. Proposta enviada
                7. Juridico / aprovacao interna
                8. Fechado ganho
                9. Fechado perdido

                ## Campos obrigatorios
                Nome da conta, segmento, ICP fit, numero de prescritores, numero de usuarios operacionais, pacientes ativos estimados, dor dominante, owner do deal, buyer economico, proximo passo, data do proximo passo, probabilidade, plano sugerido, setup estimado, ACV estimado e risco juridico / security review.

                ## Automacoes minimas e relatorios
                Criar tarefa automatica para follow-up apos discovery e apos demo. Marcar deal parado ha mais de 7 dias sem proximo passo. Exigir preenchimento de dor dominante e buyer antes de proposta enviada. Relatorios minimos: pipeline por estagio, valor ponderado, conversao discovery-demo, conversao demo-proposta, conversao proposta-ganho e tempo medio de ciclo.
                """
            ).strip(),
        },
        {
            'key': 'objection-handling-playbook',
            'title': 'Objection handling playbook',
            'tier': 'Tier 3',
            'summary': 'Scripts e talk tracks completos.',
            'pdf': True,
            'markdown': dedent(
                """
                # Objection Handling Playbook - Prescripta

                ## Ja faco isso no WhatsApp
                Resposta curta: WhatsApp resolve mensagem, nao resolve governanca de rotina.
                Talk track: O ponto nao e substituir um canal por outro. O ponto e que hoje prazo, ownership, historico e pendencia provavelmente estao espalhados. A Prescripta entra exatamente para dar uma fila observavel a uma dor que hoje fica invisivel.

                ## Nao quero mais um prontuario
                Resposta curta: A Prescripta nao quer ser prontuario.
                Talk track: Nosso recorte e estreito. A tese e organizar recorrencia prescricional, nao competir com prontuario generico. Isso reduz complexidade, acelera implantacao e melhora aderencia.

                ## E IA decidindo por mim
                Resposta curta: Nao.
                Talk track: Qualquer automacao entra para apoiar fluxo, nao para tomar decisao clinica. A autoridade final continua com o medico. Essa diferenca e central para produto, contrato e politica de IA.

                ## Tenho receio com LGPD
                Resposta curta: O receio e legitimo e precisa ser tratado com documento e processo.
                Talk track: Por isso a venda vem com contrato, DPA, politica de privacidade, politica de IA e trust center. Nao respondemos LGPD com slogan; respondemos com papeis claros, anexos e processo.
                """
            ).strip(),
        },
        {
            'key': 'homepage-copy',
            'title': 'Homepage copy',
            'tier': 'Tier 3',
            'summary': 'Copy completa da landing page.',
            'pdf': True,
            'markdown': dedent(
                f"""
                # Homepage Copy - {p['brand']}

                ## Hero
                Headline: A camada operacional da recorrencia prescricional para psiquiatria privada.
                Subheadline: Organize prazo, retorno, renovacao e confirmacao em um fluxo observavel, com supervisao humana e menos improviso entre consulta e receita.
                CTA primario: Agendar demo diagnostica
                CTA secundario: Solicitar piloto pago

                ## Problema
                O problema nao e falta de canal. E falta de processo. Em muitos consultorios e clinicas, a recorrencia prescricional ainda depende de memoria, agenda, WhatsApp e retrabalho manual.

                ## Solucao
                A {p['brand']} organiza a rotina entre termino previsto, retorno, renovacao e acompanhamento operacional. Em vez de prometer resolver tudo, entra no ponto em que a friccao realmente acontece.

                ## Confianca e CTA final
                Politica de privacidade, DPA, politica de IA com supervisao humana, trilha de governanca e separacao clara entre entrega atual e roadmap regulatorio.
                Se a sua operacao ainda depende de memoria, ela ainda nao esta pronta para escalar.
                """
            ).strip(),
        },
        {
            'key': 'revenue-projection',
            'title': 'Revenue projection spreadsheet',
            'tier': 'Tier 3',
            'summary': 'Planilha de MRR por plano, churn e break-even.',
            'xlsx': True,
        },
    ]


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    assets = base_assets()
    catalog = []
    for asset in assets:
        files = []
        markdown = asset.get('markdown')
        if markdown:
            md_path = OUTPUT_DIR / f"{asset['key']}.md"
            md_path.write_text(markdown + '\n', encoding='utf-8')
            files.append({'label': 'Markdown', 'path': f"/documents/{md_path.name}"})
            if asset.get('pdf'):
                pdf_path = OUTPUT_DIR / f"{asset['key']}.pdf"
                save_pdf(pdf_path, markdown)
                files.append({'label': 'PDF', 'path': f"/documents/{pdf_path.name}"})
            if asset.get('docx'):
                docx_path = OUTPUT_DIR / f"{asset['key']}.docx"
                save_docx(docx_path, markdown)
                files.append({'label': 'DOCX', 'path': f"/documents/{docx_path.name}"})
        if asset.get('xlsx'):
            xlsx_path = OUTPUT_DIR / f"{asset['key']}.xlsx"
            save_revenue_xlsx(xlsx_path)
            files.append({'label': 'XLSX', 'path': f"/documents/{xlsx_path.name}"})
        catalog.append(
            {
                'key': asset['key'],
                'title': asset['title'],
                'tier': asset['tier'],
                'summary': asset['summary'],
                'files': files,
            }
        )

    (OUTPUT_DIR / 'catalog.json').write_text(json.dumps(catalog, indent=2), encoding='utf-8')


if __name__ == '__main__':
    main()
