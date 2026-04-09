from __future__ import annotations

from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / 'public' / 'documents'

PDF_TO_MD = {
    'termos-de-uso.pdf': 'termos-de-uso.md',
    'politica-de-privacidade.pdf': 'politica-de-privacidade.md',
    'dpa.pdf': 'dpa.md',
    'politica-de-ia.pdf': 'politica-de-ia.md',
    'proposta-comercial.pdf': 'proposta-comercial.md',
    'one-pager.pdf': 'one-pager.md',
    'contrato-saas-msa.pdf': 'contrato-saas-msa.md',
    'playbook-onboarding.pdf': 'playbook-onboarding.md',
    'discovery-call-script.pdf': 'discovery-call-script.md',
    'welcome-email-sequence.pdf': 'welcome-email-sequence.md',
    'trust-center-content.pdf': 'trust-center-content.md',
    'crm-setup-spec.pdf': 'crm-setup-spec.md',
    'objection-handling-playbook.pdf': 'objection-handling-playbook.md',
    'homepage-copy.pdf': 'homepage-copy.md',
}


def extract_text(pdf_path: Path) -> str:
    reader = PdfReader(str(pdf_path))
    chunks = []
    for page in reader.pages:
        text = page.extract_text() or ''
        chunks.append(text)
    content = '\n'.join(chunks)
    content = content.replace('\x7f', '- ')
    content = content.replace('\u2022', '- ')
    return '\n'.join(line.rstrip() for line in content.splitlines()).strip() + '\n'


def to_docx(text: str, output_path: Path):
    document = Document()
    for line in text.splitlines():
        clean = line.strip()
        if not clean:
            document.add_paragraph('')
            continue
        if clean.startswith('- '):
            document.add_paragraph(clean[2:], style='List Bullet')
            continue
        document.add_paragraph(clean)
    document.save(output_path)


def main():
    for pdf_name, md_name in PDF_TO_MD.items():
        text = extract_text(DOCS / pdf_name)
        (DOCS / md_name).write_text(text, encoding='utf-8')

    proposal_text = (DOCS / 'proposta-comercial.md').read_text(encoding='utf-8')
    to_docx(proposal_text, DOCS / 'proposta-comercial.docx')


if __name__ == '__main__':
    main()
